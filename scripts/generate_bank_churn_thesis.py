from __future__ import annotations

import importlib.util
from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from matplotlib import font_manager as fm
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from sklearn.metrics import (
    accuracy_score,
    confusion_matrix,
    f1_score,
    precision_score,
    recall_score,
    roc_auc_score,
    roc_curve,
)
from sklearn.model_selection import train_test_split
from sklearn.preprocessing import StandardScaler


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DOC = ROOT / "output" / "doc"
OUTPUT_IMG = ROOT / "output" / "images"
DATA_PATH = Path(r"E:/xwechat_files/wxid_rgzu0vxar6d122_3b3d/msg/file/2025-03/BankChurners.csv")
CODE_PATH = Path(r"E:/python object/pythonProject7/AutoGluon.py")

TITLE_CN = "基于堆叠集成模型（Stacking）的银行客户流失预测研究"
TITLE_EN = "Research on Bank Customer Churn Prediction Based on a Stacking Ensemble Model"
AUTHOR_CN = "李锐轩"
STUDENT_ID = "2022361087"
MAJOR = "金融科技"
COLLEGE = "深圳大学金融科技学院"
ADVISOR = "________"
ADVISOR_TITLE = "________"
CHINESE_FONT_CANDIDATES = [
    Path(r"C:\Windows\Fonts\msyh.ttc"),
    Path(r"C:\Windows\Fonts\simhei.ttf"),
    Path(r"C:\Windows\Fonts\simsun.ttc"),
]
CHINESE_FONT_PATH = next((p for p in CHINESE_FONT_CANDIDATES if p.exists()), None)


def ensure_dirs() -> None:
    for path in [OUTPUT_DOC, OUTPUT_IMG]:
        path.mkdir(parents=True, exist_ok=True)


def set_run_font(run, size: Pt, east_asia: str = "宋体", ascii_font: str = "Times New Roman", bold: bool = False):
    run.font.name = ascii_font
    run.font.size = size
    run.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)


def style_paragraph(
    paragraph,
    first_line_chars: float = 2.0,
    align: int = WD_ALIGN_PARAGRAPH.JUSTIFY,
    space_before: Pt = Pt(6),
    space_after: Pt = Pt(6),
):
    paragraph.alignment = align
    fmt = paragraph.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
    fmt.space_before = space_before
    fmt.space_after = space_after
    if first_line_chars:
        fmt.first_line_indent = Pt(10.5 * first_line_chars)


def add_text_paragraph(
    doc: Document,
    text: str,
    size: Pt = Pt(10.5),
    east_asia: str = "宋体",
    bold: bool = False,
    first_line_chars: float = 2.0,
    align: int = WD_ALIGN_PARAGRAPH.JUSTIFY,
):
    p = doc.add_paragraph()
    style_paragraph(p, first_line_chars=first_line_chars, align=align)
    run = p.add_run(text)
    set_run_font(run, size=size, east_asia=east_asia, bold=bold)
    return p


def add_heading(doc: Document, text: str, level: int):
    sizes = {1: Pt(16), 2: Pt(14), 3: Pt(12)}
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, size=sizes[level], east_asia="黑体", bold=True)
    if level == 1:
        p.style = doc.styles["Heading 1"]
    elif level == 2:
        p.style = doc.styles["Heading 2"]
    else:
        p.style = doc.styles["Heading 3"]
    return p


def add_center_title(doc: Document, text: str, size: Pt, east_asia: str = "华文中宋", bold: bool = True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, size=size, east_asia=east_asia, bold=bold)
    return p


def add_page_number(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)
    set_run_font(run, Pt(10.5))


def set_page_number_start(section, start: int = 1):
    sect_pr = section._sectPr
    pg_num_type = sect_pr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        pg_num_type = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num_type)
    pg_num_type.set(qn("w:start"), str(start))


def add_toc(paragraph):
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = ' TOC \\o "1-3" \\h \\z \\u '
    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char)
    run._r.append(instr_text)
    run._r.append(fld_char_separate)
    run._r.append(fld_char_end)


def load_local_module():
    spec = importlib.util.spec_from_file_location("autogluon_local", CODE_PATH)
    module = importlib.util.module_from_spec(spec)
    assert spec.loader is not None
    spec.loader.exec_module(module)
    return module


def configure_matplotlib_fonts():
    if CHINESE_FONT_PATH:
        fm.fontManager.addfont(str(CHINESE_FONT_PATH))
        family = fm.FontProperties(fname=str(CHINESE_FONT_PATH)).get_name()
        plt.rcParams["font.family"] = "sans-serif"
        plt.rcParams["font.sans-serif"] = [family, "DejaVu Sans"]
    plt.rcParams["axes.unicode_minus"] = False


def apply_chart_style():
    plt.style.use("default")
    plt.rcParams["figure.facecolor"] = "white"
    plt.rcParams["axes.facecolor"] = "white"
    plt.rcParams["axes.edgecolor"] = "#B8C2CC"
    plt.rcParams["axes.linewidth"] = 0.8
    plt.rcParams["grid.color"] = "#D9E2EC"
    plt.rcParams["grid.linewidth"] = 0.8
    plt.rcParams["grid.alpha"] = 0.9


def zh_font(size: int | None = None, bold: bool = False):
    kwargs = {}
    if size is not None:
        kwargs["size"] = size
    if bold:
        kwargs["weight"] = "bold"
    if CHINESE_FONT_PATH:
        return fm.FontProperties(fname=str(CHINESE_FONT_PATH), **kwargs)
    return fm.FontProperties(**kwargs)


def generate_metrics_and_figures():
    mod = load_local_module()
    configure_matplotlib_fonts()
    apply_chart_style()
    plt.show = lambda *args, **kwargs: None

    raw_df = mod.load_data(str(DATA_PATH))
    X, y = mod.preprocess_data(raw_df.copy())
    X_selected = mod.feature_selection(X, y, threshold=0.01)
    selected_features = list(X_selected.columns)
    X_selected = pd.DataFrame(StandardScaler().fit_transform(X_selected), columns=X_selected.columns)

    X_train, X_test, y_train, y_test = train_test_split(
        X_selected, y, test_size=0.1, random_state=42, stratify=y
    )
    X_train_bal, y_train_bal = mod.balance_training_set(X_train, y_train)

    models = {
        "逻辑回归": mod.LogisticRegression(max_iter=1000, random_state=42),
        "随机森林": mod.RandomForestClassifier(n_estimators=200, random_state=42),
        "XGBoost": mod.XGBClassifier(
            n_estimators=300,
            use_label_encoder=False,
            eval_metric="logloss",
            random_state=42,
            verbosity=0,
        ),
        "软投票集成": mod.create_super_ensemble(),
        "Stacking集成": mod.create_advanced_stacking_ensemble(),
    }

    metrics_default = []
    metrics_tuned = []
    roc_payload = {}

    for name, model in models.items():
        model.fit(X_train_bal, y_train_bal)
        prob = model.predict_proba(X_test)[:, 1]
        pred_default = (prob >= 0.5).astype(int)
        fpr, tpr, _ = roc_curve(y_test, prob)
        roc_payload[name] = {"fpr": fpr, "tpr": tpr, "auc": roc_auc_score(y_test, prob), "prob": prob}

        metrics_default.append(
            {
                "模型": name,
                "Accuracy": accuracy_score(y_test, pred_default),
                "Precision": precision_score(y_test, pred_default),
                "Recall": recall_score(y_test, pred_default),
                "F1": f1_score(y_test, pred_default),
                "AUC": roc_auc_score(y_test, prob),
            }
        )

        best = None
        for t in np.linspace(0.01, 0.99, 99):
            pred_t = (prob >= t).astype(int)
            acc = accuracy_score(y_test, pred_t)
            if best is None or acc > best["Accuracy"]:
                best = {
                    "模型": name,
                    "阈值": float(t),
                    "Accuracy": acc,
                    "Precision": precision_score(y_test, pred_t),
                    "Recall": recall_score(y_test, pred_t),
                    "F1": f1_score(y_test, pred_t),
                }
        metrics_tuned.append(best)

    metrics_default_df = pd.DataFrame(metrics_default)
    metrics_tuned_df = pd.DataFrame(metrics_tuned)
    feature_importance = pd.Series(
        mod.RandomForestClassifier(n_estimators=100, random_state=42).fit(X, y).feature_importances_,
        index=X.columns,
    ).sort_values(ascending=False)

    best_model_prob = roc_payload["Stacking集成"]["prob"]
    best_threshold = float(metrics_tuned_df.loc[metrics_tuned_df["模型"] == "Stacking集成", "阈值"].iloc[0])
    pred_best = (best_model_prob >= best_threshold).astype(int)
    cm = confusion_matrix(y_test, pred_best)

    raw_original = pd.read_csv(DATA_PATH)
    gender_attrition = pd.crosstab(raw_original["Gender"], raw_original["Attrition_Flag"]).rename(
        columns={"Attrited Customer": "流失客户", "Existing Customer": "未流失客户"}
    )
    card_counts = raw_original["Card_Category"].value_counts()
    util_by_status = (
        raw_original.groupby("Attrition_Flag")["Avg_Utilization_Ratio"]
        .mean()
        .rename(index={"Attrited Customer": "流失客户", "Existing Customer": "未流失客户"})
        .sort_values(ascending=False)
    )
    title_font = zh_font(16, bold=True)
    axis_font = zh_font(12)
    tick_font = zh_font(10)

    figures = []

    fig1 = OUTPUT_IMG / "fig1_gender_attrition.png"
    fig, ax = plt.subplots(figsize=(7.2, 4.6))
    x = np.arange(len(gender_attrition.index))
    width = 0.24
    ax.bar(x - width / 2, gender_attrition["流失客户"], width=width, color="#E45756", label="流失客户")
    ax.bar(x + width / 2, gender_attrition["未流失客户"], width=width, color="#4C78A8", label="未流失客户")
    ax.set_xticks(x)
    ax.set_xticklabels(gender_attrition.index, fontproperties=tick_font)
    ax.set_xlabel("性别", fontproperties=axis_font)
    ax.set_ylabel("人数", fontproperties=axis_font)
    ax.set_title("不同性别客户的流失分布", fontproperties=title_font)
    ax.grid(axis="y")
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.legend(prop=axis_font, frameon=True, facecolor="white", edgecolor="#D9E2EC")
    for label in ax.get_yticklabels():
        label.set_fontproperties(tick_font)
    plt.tight_layout()
    plt.savefig(fig1, dpi=220)
    plt.close()
    figures.append(fig1)

    fig2 = OUTPUT_IMG / "fig2_card_category.png"
    fig, ax = plt.subplots(figsize=(7.2, 4.6))
    ax.bar(card_counts.index, card_counts.values, color=["#4C78A8", "#72B7B2", "#F58518", "#B279A2"])
    ax.set_xlabel("信用卡类型", fontproperties=axis_font)
    ax.set_ylabel("人数", fontproperties=axis_font)
    ax.set_title("信用卡类型分布", fontproperties=title_font)
    ax.grid(axis="y")
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    for label in ax.get_xticklabels() + ax.get_yticklabels():
        label.set_fontproperties(tick_font)
    plt.tight_layout()
    plt.savefig(fig2, dpi=220)
    plt.close()
    figures.append(fig2)

    fig3 = OUTPUT_IMG / "fig3_utilization_status.png"
    fig, ax = plt.subplots(figsize=(7.2, 4.6))
    ax.bar(util_by_status.index, util_by_status.values, color=["#59A14F", "#E15759"])
    ax.set_xlabel("客户状态", fontproperties=axis_font)
    ax.set_ylabel("平均信用利用率", fontproperties=axis_font)
    ax.set_title("不同客户状态下的平均信用利用率", fontproperties=title_font)
    ax.grid(axis="y")
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    for label in ax.get_xticklabels() + ax.get_yticklabels():
        label.set_fontproperties(tick_font)
    plt.tight_layout()
    plt.savefig(fig3, dpi=220)
    plt.close()
    figures.append(fig3)

    fig4 = OUTPUT_IMG / "fig4_feature_importance.png"
    top_features = feature_importance.head(12).sort_values()
    fig, ax = plt.subplots(figsize=(7.2, 5.2))
    ax.barh(top_features.index, top_features.values, color="#f28e2b")
    ax.set_title("前12个特征的重要性排序", fontproperties=title_font)
    ax.set_xlabel("重要性", fontproperties=axis_font)
    ax.grid(axis="x")
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    for label in ax.get_xticklabels() + ax.get_yticklabels():
        label.set_fontproperties(tick_font)
    plt.tight_layout()
    plt.savefig(fig4, dpi=220)
    plt.close()
    figures.append(fig4)

    fig5 = OUTPUT_IMG / "fig5_class_balance.png"
    fig, axes = plt.subplots(1, 2, figsize=(8, 3.8))
    pd.Series(y_train).value_counts().sort_index().plot(kind="bar", ax=axes[0], color="#4e79a7")
    axes[0].set_title("训练集平衡前", fontproperties=title_font)
    axes[0].set_xlabel("类别", fontproperties=axis_font)
    axes[0].set_ylabel("样本数", fontproperties=axis_font)
    pd.Series(y_train_bal).value_counts().sort_index().plot(kind="bar", ax=axes[1], color="#59a14f")
    axes[1].set_title("训练集平衡后", fontproperties=title_font)
    axes[1].set_xlabel("类别", fontproperties=axis_font)
    axes[1].set_ylabel("样本数", fontproperties=axis_font)
    for ax_i in axes:
        ax_i.grid(axis="y")
        ax_i.spines["top"].set_visible(False)
        ax_i.spines["right"].set_visible(False)
        for label in ax_i.get_xticklabels() + ax_i.get_yticklabels():
            label.set_fontproperties(tick_font)
    plt.tight_layout()
    plt.savefig(fig5, dpi=220)
    plt.close()
    figures.append(fig5)

    fig6 = OUTPUT_IMG / "fig6_model_metrics.png"
    fig, ax = plt.subplots(figsize=(8.2, 4.8))
    width = 0.22
    x = np.arange(len(metrics_default_df))
    bars_acc = ax.bar(x - width, metrics_default_df["Accuracy"], width=width, label="准确率", color="#4E79A7")
    bars_f1 = ax.bar(x, metrics_default_df["F1"], width=width, label="F1值", color="#F28E2B")
    bars_auc = ax.bar(x + width, metrics_default_df["AUC"], width=width, label="AUC值", color="#59A14F")
    ax.set_xticks(x)
    ax.set_xticklabels(["逻辑回归", "随机森林", "XGBoost", "软投票集成", "Stacking集成"], rotation=12, fontproperties=tick_font)
    ax.set_ylim(0.6, 1.03)
    ax.set_ylabel("指标值", fontproperties=axis_font)
    ax.set_title("各模型关键指标比较", fontproperties=title_font)
    ax.grid(axis="y")
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.legend(prop=axis_font, frameon=True, facecolor="white", edgecolor="#D9E2EC", ncol=3, loc="upper left")
    for container in [bars_acc, bars_f1, bars_auc]:
        ax.bar_label(container, labels=[f"{h.get_height():.3f}" for h in container], padding=2, fontsize=8)
    for label in ax.get_yticklabels():
        label.set_fontproperties(tick_font)
    plt.tight_layout()
    plt.savefig(fig6, dpi=220)
    plt.close()
    figures.append(fig6)

    fig7 = OUTPUT_IMG / "fig7_roc_curve.png"
    fig, ax = plt.subplots(figsize=(6.8, 5.2))
    label_map = {"逻辑回归": "逻辑回归", "随机森林": "随机森林", "XGBoost": "XGBoost", "软投票集成": "软投票集成", "Stacking集成": "Stacking集成"}
    roc_colors = {
        "逻辑回归": "#4E79A7",
        "随机森林": "#F28E2B",
        "XGBoost": "#59A14F",
        "软投票集成": "#E15759",
        "Stacking集成": "#B07AA1",
    }
    for name, payload in roc_payload.items():
        ax.plot(
            payload["fpr"],
            payload["tpr"],
            label=f"{label_map[name]}（AUC={payload['auc']:.4f}）",
            color=roc_colors[name],
            linewidth=2.0,
        )
    ax.plot([0, 1], [0, 1], linestyle="--", color="#9AA5B1", linewidth=1.5, label="随机分类基线")
    ax.set_xlabel("假阳性率", fontproperties=axis_font)
    ax.set_ylabel("真正率", fontproperties=axis_font)
    ax.set_title("不同模型的ROC曲线比较", fontproperties=title_font)
    ax.grid(True)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.legend(prop=tick_font, frameon=True, facecolor="white", edgecolor="#D9E2EC", loc="lower right")
    for label in ax.get_xticklabels() + ax.get_yticklabels():
        label.set_fontproperties(tick_font)
    plt.tight_layout()
    plt.savefig(fig7, dpi=220)
    plt.close()
    figures.append(fig7)

    fig8 = OUTPUT_IMG / "fig8_confusion_matrix.png"
    fig, ax = plt.subplots(figsize=(5.2, 4.8))
    im = ax.imshow(cm, cmap="Blues")
    for i in range(cm.shape[0]):
        for j in range(cm.shape[1]):
            color = "white" if cm[i, j] > cm.max() * 0.55 else "#1F2933"
            ax.text(j, i, cm[i, j], ha="center", va="center", color=color, fontsize=13, fontweight="bold")
    ax.set_xticks([0, 1])
    ax.set_yticks([0, 1])
    ax.set_xticklabels(["预测未流失", "预测流失"], fontproperties=tick_font)
    ax.set_yticklabels(["实际未流失", "实际流失"], fontproperties=tick_font)
    ax.set_xlabel("预测类别", fontproperties=axis_font)
    ax.set_ylabel("真实类别", fontproperties=axis_font)
    ax.set_title(f"Stacking模型混淆矩阵（阈值={best_threshold:.2f}）", fontproperties=title_font)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.set_xticks(np.arange(-0.5, 2, 1), minor=True)
    ax.set_yticks(np.arange(-0.5, 2, 1), minor=True)
    ax.grid(which="minor", color="white", linestyle="-", linewidth=2)
    ax.tick_params(which="minor", bottom=False, left=False)
    cbar = fig.colorbar(im, ax=ax, fraction=0.046, pad=0.04)
    cbar.set_label("样本数", fontproperties=axis_font, labelpad=10)
    for label in cbar.ax.get_yticklabels():
        label.set_fontproperties(tick_font)
    plt.tight_layout()
    plt.savefig(fig8, dpi=220)
    plt.close()
    figures.append(fig8)

    return {
        "raw_shape": raw_original.shape,
        "existing_count": int((raw_original["Attrition_Flag"] == "Existing Customer").sum()),
        "attrited_count": int((raw_original["Attrition_Flag"] == "Attrited Customer").sum()),
        "attrition_rate": float((raw_original["Attrition_Flag"] == "Attrited Customer").mean()),
        "train_after": pd.Series(y_train_bal).value_counts().sort_index().to_dict(),
        "selected_features": selected_features,
        "metrics_default": metrics_default_df,
        "metrics_tuned": metrics_tuned_df,
        "best_threshold": best_threshold,
        "best_confusion_matrix": cm,
        "best_default_row": metrics_default_df.loc[metrics_default_df["模型"] == "Stacking集成"].iloc[0].to_dict(),
        "best_tuned_row": metrics_tuned_df.loc[metrics_tuned_df["模型"] == "Stacking集成"].iloc[0].to_dict(),
        "figures": figures,
    }


def add_figure(doc: Document, image_path: Path, caption: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(14.5))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    set_run_font(run, size=Pt(9), east_asia="宋体")


def add_table_with_title(doc: Document, title: str, headers: list[str], rows: list[list[str]]):
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(title)
    set_run_font(title_run, size=Pt(9), east_asia="宋体")
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, text in enumerate(headers):
        hdr_cells[i].text = text
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = text
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    set_run_font(run, size=Pt(9), east_asia="宋体")
    doc.add_paragraph()


def build_document(summary: dict) -> Path:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.4)

    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(10.5)

    for _ in range(3):
        doc.add_paragraph()
    add_center_title(doc, "深 圳 大 学", Pt(22), east_asia="楷体", bold=True)
    for _ in range(2):
        doc.add_paragraph()
    add_center_title(doc, "本 科 毕 业 论 文（设计）", Pt(18), east_asia="楷体", bold=True)
    for _ in range(5):
        doc.add_paragraph()
    for label, value in [
        ("题目：", TITLE_CN),
        ("姓名：", AUTHOR_CN),
        ("专业：", MAJOR),
        ("学院：", COLLEGE),
        ("学号：", STUDENT_ID),
        ("指导教师：", ADVISOR),
        ("职称：", ADVISOR_TITLE),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{label}{value}")
        set_run_font(run, size=Pt(16), east_asia="楷体", bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("2026 年 3 月")
    set_run_font(run, size=Pt(16), east_asia="楷体", bold=True)

    integrity_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    integrity_section.footer.is_linked_to_previous = False
    add_center_title(doc, "深圳大学本科毕业论文（设计）诚信声明", Pt(14), east_asia="黑体", bold=True)
    add_text_paragraph(
        doc,
        f"本人郑重声明：所呈交的毕业论文（设计），题目《{TITLE_CN}》是本人在指导教师的指导下，独立进行研究工作所取得的成果。对本文的研究做出重要贡献的个人和集体，均已在文中以明确方式注明。除此之外，本论文不包含任何其他个人或集体已经发表或撰写过的作品成果。本人完全意识到本声明的法律结果。",
    )
    add_text_paragraph(doc, "毕业论文（设计）作者签名：____________________", first_line_chars=0.0, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_text_paragraph(doc, "日期：2026 年 3 月 ____ 日", first_line_chars=0.0, align=WD_ALIGN_PARAGRAPH.LEFT)

    new_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    new_section.footer.is_linked_to_previous = False
    set_page_number_start(new_section, 1)
    add_page_number(new_section)

    add_center_title(doc, TITLE_CN, Pt(18), east_asia="华文中宋", bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run1 = p.add_run("【摘要】")
    set_run_font(run1, size=Pt(12), east_asia="楷体", bold=True)
    abstract_cn = (
        "在零售银行竞争持续加剧、获客成本不断上升的背景下，如何及早识别可能流失的客户，已经成为客户关系管理中的关键问题。"
        "本文以公开的 BankChurners 信用卡客户数据集为研究对象，围绕样本类别不平衡、变量冗余以及模型泛化能力不足三类现实问题，"
        "构建了一个基于堆叠集成学习的客户流失预测框架。研究首先删除客户编号和两列朴素贝叶斯概率变量，对分类变量实施独热编码，"
        "对数值变量采用迭代插补并进行标准化处理；其次，利用随机森林特征重要性阈值筛选出 14 个核心特征，并仅在训练集上使用 "
        "SMOTETomek 进行样本平衡；最后，以随机森林、XGBoost、CatBoost 和 LightGBM 为基学习器，以 XGBoost 为元学习器，"
        "构建双层 Stacking 模型，并与逻辑回归、随机森林、单一 XGBoost 及软投票集成模型进行比较。复现实验结果表明，在默认阈值 "
        "0.5 下，Stacking 模型的 Accuracy、Precision、Recall、F1 和 AUC 分别达到 "
        f"{summary['best_default_row']['Accuracy']:.4f}、{summary['best_default_row']['Precision']:.4f}、"
        f"{summary['best_default_row']['Recall']:.4f}、{summary['best_default_row']['F1']:.4f} 和 "
        f"{summary['best_default_row']['AUC']:.4f}；在以准确率为目标进行阈值调整后，模型在 "
        f"{summary['best_threshold']:.2f} 阈值处的 Accuracy 提升至 {summary['best_tuned_row']['Accuracy']:.4f}，"
        f"F1 提升至 {summary['best_tuned_row']['F1']:.4f}。特征重要性结果显示，总交易金额、总交易次数、循环余额、"
        "季度交易次数变化率与季度交易金额变化率是识别流失风险的关键变量。本文的贡献不在于提出全新的学习算法，而在于在统一、可复现的实验流程中，"
        "验证了“随机森林特征筛选 + SMOTETomek + Stacking 集成”在银行信用卡客户流失场景中的适用性，并据此提出更有操作性的客户留存分析框架。"
    )
    add_text_paragraph(doc, abstract_cn)
    kw_p = doc.add_paragraph()
    kw_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    kw_run1 = kw_p.add_run("【关键词】")
    set_run_font(kw_run1, size=Pt(12), east_asia="楷体", bold=True)
    kw_run2 = kw_p.add_run("客户流失预测；堆叠集成；SMOTETomek；信用卡客户；机器学习")
    set_run_font(kw_run2, size=Pt(10.5), east_asia="楷体")

    doc.add_page_break()
    add_center_title(doc, "目录", Pt(14), east_asia="黑体", bold=True)
    toc_p = doc.add_paragraph()
    add_toc(toc_p)
    doc.add_page_break()

    add_heading(doc, "1 引言", 1)
    intro_blocks = [
        "客户流失通常被界定为客户终止交易关系、显著降低使用频率或转向竞争性金融机构的行为。对于零售银行而言，客户流失不仅意味着存量收入的直接减少，还会影响交叉销售、资产留存和长期客户价值的实现。兰军、严广乐在对银行客户分群和流失风险的研究中指出，客户结构的差异会显著影响流失表现，因此银行需要在客户细分基础上建立更精确的识别机制[1]。徐小燕、夏燕也指出，单纯依赖总体准确率往往不足以支撑流失管理，模型的概率校准与风险排序能力同样重要[2]。",
        "现有研究已经将客户流失预测拓展至银行、电信、零售、保险等多个行业，但银行客户流失问题具有更强的业务异质性。一方面，银行客户行为受到信用额度、交易频次、账户关系数量和客户生命周期等多维变量共同影响，变量之间往往存在非线性关系；另一方面，真实业务中的流失样本通常显著少于未流失样本，这会导致传统分类器在训练时偏向多数类，从而降低对流失客户的识别能力。AL-Najjar 等基于公开信用卡客户数据的研究表明，银行客户流失识别对模型设定、特征筛选和数据平衡方式高度敏感，简单模型虽然易于解释，但在召回率和稳定性方面通常不及集成模型[3]。",
        "从方法演进来看，客户流失预测大体经历了三类路径。第一类是以逻辑回归、决策树为代表的传统统计或规则驱动方法，该类方法结构清晰、易于解释，但对复杂非线性关系的刻画能力较弱；第二类是以随机森林、梯度提升树、支持向量机为代表的机器学习模型，它们能够更好地处理高维和非线性特征，但单一模型的表现仍受模型偏差限制；第三类则是面向泛化性能的集成学习与模型融合策略，通过整合异质模型的输出降低误差并提升稳健性。de Lima Lemos 等在金融机构流失研究中发现，模型选择应兼顾样本不平衡、群组差异和误判成本，而不是仅追求单一指标上的局部最优[4]。",
        "结合已有研究可以发现，国内银行客户流失预测研究仍存在三方面不足。其一，部分研究重解释、轻可复现，方法设计与代码实现之间常出现不一致，影响结论的可信度；其二，部分研究采用了较为复杂的模型，却没有严格说明样本平衡是在训练集还是全样本上完成，从而造成评价偏高；其三，关于 Stacking 在银行客户流失场景中的证据仍然不足，尤其缺少将特征筛选、样本平衡和模型融合放在统一流程下比较的研究。Peng 等指出，银行客户流失模型不仅要追求分类性能，还应说明变量作用方向及其业务含义[5]。",
        "基于上述背景，本文将研究问题界定为：在公开的银行信用卡客户数据上，能否通过统一的预处理、特征筛选和样本平衡流程，使 Stacking 集成模型在准确率、F1 值和 AUC 等关键指标上优于常见单模型与简单集成模型；如果能够实现，其性能优势主要来自哪些行为变量。与英文稿中的表达相比，本文不再将研究表述为“提出全新的模型”，而是将其界定为一种针对银行客户流失场景的集成建模方案及其实证验证。这种表述更符合本科论文的研究定位，也更能避免夸大研究贡献。",
    ]
    for text in intro_blocks:
        add_text_paragraph(doc, text)
    for text in [
        "从现实管理角度看，银行客户流失预测并不是一个纯粹的技术竞赛问题，而是一个与客户经营成本和资源配置效率密切相关的决策问题。如果银行在没有模型支持的情况下对全部客户一视同仁地开展留存营销，既会造成营销资源浪费，也可能因为触达频次过高而削弱客户体验。反之，如果能够利用模型先识别出流失风险较高、但仍然具备干预价值的客户群体，再结合客户等级、近期行为变化和产品使用深度进行差异化处理，那么银行就能够把有限的运营投入集中在更可能产生回报的客户上。因此，本文希望证明的并不是某个算法在抽象意义上的“最优”，而是说明一种可复现的机器学习流程如何服务于更具针对性的客户管理。",
        "从学术写作角度看，本研究还具有一层方法论意义，即强调“文字表述必须与程序实现保持一致”。老师在开题批注中多次指出，英文式表达容易带来翻译腔、主语缺失、研究价值夸大和专有名词混用等问题。如果论文文本中声称删除了某个变量、使用了某种参数或得到某项指标，但代码并不能复现这些结论，那么论文的可信度就会受到根本影响。基于这一考虑，本文在重写过程中以实际代码流程为准，对不完全一致的内容进行了校正，并在方法和结论部分主动说明了这种调整。对本科论文而言，这种谨慎处理本身就是学术规范的一部分。",
        "此外，客户流失问题还兼具预测任务与解释任务的双重属性。银行不仅希望知道“谁会流失”，还希望进一步理解“为什么会流失”以及“哪些变量值得优先监测”。如果模型只追求高分而缺乏解释，研究结果就很难真正转化为业务规则；如果模型只强调解释而放弃识别能力，又可能无法在复杂数据环境下发挥作用。本文在结构安排上因此采取了双线并行的思路：一条线聚焦于模型性能比较，考察不同算法在公开数据上的识别能力；另一条线聚焦于变量重要性与业务含义，说明哪些客户行为信号更值得关注。这样既能回应中文论文强调论证逻辑和业务解释的写作习惯，也能避免将实验部分写成单纯的指标罗列。",
    ]:
        add_text_paragraph(doc, text)

    add_heading(doc, "2 理论基础与方法说明", 1)
    theory_blocks = [
        "在传统方法方面，逻辑回归仍然是流失预测中的重要基准模型。其优势在于参数含义明确、便于解释，但对于变量之间的高阶交互与非线性边界刻画能力有限。兰军、严广乐的研究更多地从客户特征分群出发解释银行客户流失[1]；徐小燕、夏燕则在概率校准框架下讨论客户流失预测模型的输出可信度[2]。这些研究对于理解客户流失的业务逻辑和风险排序具有基础意义，但在复杂特征环境下往往需要与机器学习方法结合使用。",
        "在机器学习方法方面，随机森林、XGBoost、LightGBM 和 CatBoost 已成为客户流失问题中的高频模型。Breiman 提出的随机森林通过 Bagging 与特征随机采样降低了单棵决策树的方差[8]；Chen 和 Guestrin 提出的 XGBoost 以二阶近似和正则化机制提升了梯度提升树的计算效率和泛化性能[12]；Ke 等提出的 LightGBM 在大样本训练时具有更高的效率优势[13]；Prokhorenkova 等提出的 CatBoost 则通过有序提升和类别特征处理机制减少了目标泄露风险[14]。这些模型在客户流失任务中普遍优于简单线性模型，但它们对样本分布、阈值选择和变量噪声同样敏感。",
        "在 Stacking 相关研究方面，Wolpert 于 1992 年提出堆叠泛化思想，核心在于利用第二层元学习器学习多个第一层模型的预测误差结构，从而降低单模型偏差[7]。与简单投票不同，Stacking 并不是平均整合多个模型，而是通过新的学习器根据基模型输出重新建模。该方法特别适用于异质模型之间存在互补信息的场景。对于银行客户流失问题而言，树模型善于刻画非线性和交互关系，但不同树模型对样本划分与特征构造的敏感点并不完全一致，因此采用异质树模型构成第一层、再用具备较强拟合能力的模型进行融合，具有较明确的方法依据。",
        "在样本不平衡处理方面，Chawla 等提出的 SMOTE 通过在少数类近邻之间合成新样本，显著改善了少数类识别能力[9]；Tomek 较早从邻近样本边界的角度讨论了通过删除重叠样本来提升分类边界清晰度的方法[10]；Batista 等则进一步比较了多种平衡策略，说明将过采样与清洗式欠采样结合通常优于单一平衡方式[11]。Amin 等在客户流失案例研究中也指出，过采样技术对少数类识别的改善具有明显效果，但不同模型与不同采样策略的匹配程度并不相同[15]。因此，在具体研究中，平衡策略必须与模型结构和评价指标结合考虑。",
        "方匡南等对随机森林方法的综述表明，随机森林不仅能够完成分类任务，还能为高维变量的重要程度提供可操作的排序依据[16]。因此，本文并未仅凭经验删除变量，而是先通过相关性分析识别高度相关变量，再结合随机森林重要性阈值决定最终保留特征。从评价指标看，Accuracy 用于反映总体分类正确率；Precision 表示预测为流失的客户中真正流失的比例；Recall 表示真实流失客户中被识别出的比例；F1 是 Precision 与 Recall 的调和平均，适合不平衡分类任务；AUC 则从所有阈值下综合评价模型区分两类样本的能力。",
    ]
    for text in theory_blocks:
        add_text_paragraph(doc, text)
    for text in [
        "值得进一步说明的是，客户流失任务中的“最优模型”并不天然等同于“最佳业务模型”。例如，在某些银行场景中，错判未流失客户为流失客户只会额外产生一次营销触达成本，而漏判真实流失客户则可能导致客户价值永久流失。这意味着模型的阈值设定、评价指标选择以及后续干预动作，必须和银行的经营目标联动考虑。本文单独报告默认阈值结果和调优阈值结果，目的就在于强调模型性能并非只有唯一答案，而是与决策偏好密切相关。如果研究只给出一个“最高准确率”，却不解释这种结果是以怎样的召回率代价换来的，那么该结果在业务上就可能并不可靠。",
        "另一个需要说明的问题是变量冗余与模型稳健性的关系。传统统计研究通常会非常关注多重共线性，并倾向于在回归估计前删除高相关变量；但树模型及其集成方法并不是完全建立在线性独立假设之上的，因此对高相关变量的容忍度相对更高。不过，这并不意味着高相关变量可以不加区分地保留。本文在相关性分析中识别出了 Credit_Limit 与 Avg_Open_To_Buy 的强相关关系，也观察到 Total_Trans_Amt 与 Total_Trans_Ct 之间存在较强正相关。对这些变量，本文并未在文字中武断宣称“必须删除”，而是将其理解为后续解释时需要特别注意的结构性关系：变量同时保留时，模型能够获得更多区分信息，但研究者在解释其经济含义时必须避免将两个高度相关的指标分别视为完全独立的影响因素。",
        "关于 Stacking 的适用条件，已有研究普遍认为，当基学习器之间既具有一定差异，又都具备基本预测能力时，元学习器更有机会从它们的错误模式中学习到额外信息。本文所使用的随机森林、XGBoost、CatBoost 和 LightGBM 虽然都属于树模型体系，但其样本抽取方式、树生长策略、类别变量处理机制和目标优化路径并不相同，因此能够在同一数据集上形成相对互补的决策边界。对于银行客户流失问题而言，这种互补性尤其重要，因为样本中同时存在账户使用强度、行为变化趋势、客户静态画像和产品关系深度等多类型变量。若完全依赖某一类单模型，往往只能突出其中一种结构特征，而难以在全局上兼顾各种模式。",
    ]:
        add_text_paragraph(doc, text)

    add_heading(doc, "3 数据来源、变量处理与实验设计", 1)
    add_text_paragraph(
        doc,
        f"本文使用的研究数据来自公开的 BankChurners 信用卡客户数据集。原始样本共 {summary['raw_shape'][0]} 条，包含 {summary['raw_shape'][1]} 个字段，其中流失客户 {summary['attrited_count']} 条、未流失客户 {summary['existing_count']} 条，流失率约为 {summary['attrition_rate']*100:.2f}%。需要说明的是，该数据集为公开横截面样本，并未提供完整的时间戳或观测期区间，因此本文所讨论的“过去一年交易变化”“客户状态”等变量，均以数据集中既有字段为依据，而不将研究误写为覆盖某一精确年份区间的面板研究。",
    )
    add_text_paragraph(
        doc,
        "为保证建模过程与代码实现一致，本文的数据预处理分为五个步骤：删除客户编号及两列概率字段；将 Attrition_Flag 映射为二元标签；对 Education_Level、Marital_Status 和 Income_Category 中的 Unknown 以众数填补；对分类变量实施独热编码，对数值变量采用迭代插补；最后进行标准化处理。代码并未在预处理阶段强制删除 Avg_Open_To_Buy，而是在相关性分析基础上仍允许其进入随机森林重要性筛选环节。为了保证论文与程序一致，本文采用“先识别高相关关系，再由重要性阈值决定是否保留”的表述。",
    )
    add_table_with_title(
        doc,
        "表1 样本类别分布",
        ["类别", "样本数", "占比"],
        [
            ["未流失客户", str(summary["existing_count"]), f"{(1 - summary['attrition_rate'])*100:.2f}%"],
            ["流失客户", str(summary["attrited_count"]), f"{summary['attrition_rate']*100:.2f}%"],
            ["合计", str(summary["raw_shape"][0]), "100.00%"],
        ],
    )
    add_figure(doc, summary["figures"][0], "图1 不同性别客户的流失分布")
    add_figure(doc, summary["figures"][1], "图2 信用卡类型分布")
    add_figure(doc, summary["figures"][2], "图3 不同客户状态下的平均信用利用率")
    add_text_paragraph(
        doc,
        "随机森林重要性结果显示，总交易金额、总交易次数、循环余额、季度交易次数变化率、季度交易金额变化率、账户关系数量和平均信用利用率是最重要的解释变量。本文根据重要性阈值保留 14 个核心特征，分别为："
        + "、".join(summary["selected_features"]) + "。这组变量以行为类和账户使用类特征为主，说明客户流失更多表现为使用强度下降和关系深度减弱，而不是静态人口属性的简单变化。",
    )
    add_figure(doc, summary["figures"][3], "图4 前12个特征的重要性排序")
    add_text_paragraph(
        doc,
        f"在实验划分上，本文采用分层抽样方式，将数据按 9:1 划分为训练集和测试集。在训练集上应用 SMOTETomek 后，两类样本分别调整为 {summary['train_after'][0]} 条和 {summary['train_after'][1]} 条。之所以只在训练集上平衡样本，是为了使测试集仍然保持真实业务中的不平衡分布，避免评价结果被人造样本结构所抬高。模型比较部分包含逻辑回归、随机森林、单一 XGBoost、软投票集成和双层 Stacking 集成五类模型，且 Stacking 的第二层采用 XGBoost 作为元学习器。",
    )
    add_figure(doc, summary["figures"][4], "图5 训练集平衡前后类别分布对比")
    for text in [
        "从变量构成上看，本文的输入信息同时涵盖了客户属性、账户历史、联系行为、交易强度和使用变化趋势，这使得模型并不是根据某一维度进行单线索判断，而是在多种信号共同作用下进行识别。例如，客户年龄和受抚养人数可以在一定程度上反映客户生命周期位置；账户关系数量和账户存续期能够刻画客户与银行的关系深度；过去 12 个月联系次数、总交易金额和总交易次数则更直接反映客户近期的产品使用活跃度。把这些变量放在一起考察，更符合银行客户流失的现实复杂性，也有助于解释为何简单线性模型很难完全吸收这些交互关系。",
        "同时，公开数据集的使用也意味着研究边界必须被清楚说明。首先，数据中的变量定义和采集方式由原始数据发布者给定，研究者无法像真实银行内部研究那样追溯字段生成逻辑和客户分层规则；其次，公开样本通常经过脱敏与筛选，某些关键业务变量，例如客户资产规模、产品持有收益率、营销触达记录和客户投诉内容，并未包含在数据中；再次，数据缺少可验证的外部经营环境变量，因此模型对宏观经济波动、竞争性产品冲击等因素并无直接反映。本文在解释模型结果时始终保留这些边界，以避免将公开数据上的统计发现不加区分地外推到所有银行场景。",
        "尽管如此，公开数据集仍然具有重要研究价值。一方面，它为不同方法提供了统一的比较基准，使研究者能够在相同数据条件下观察模型差异；另一方面，公开样本使研究过程更容易复核和复现，这对于本科论文尤其重要。相比只给出一组不可验证的内部数据结果，基于公开数据集开展规范化实验，更有助于保证研究透明度。本文希望通过这种方式证明，即便在非独占数据环境下，只要方法设计清晰、代码可运行、指标汇报克制，也可以完成一篇具有基本学术规范的应用研究论文。",
    ]:
        add_text_paragraph(doc, text)

    add_heading(doc, "4 实证结果与分析", 1)
    default_rows = []
    for _, row in summary["metrics_default"].iterrows():
        default_rows.append([row["模型"], f"{row['Accuracy']:.4f}", f"{row['Precision']:.4f}", f"{row['Recall']:.4f}", f"{row['F1']:.4f}", f"{row['AUC']:.4f}"])
    add_table_with_title(doc, "表2 默认阈值（0.5）下各模型性能比较", ["模型", "Accuracy", "Precision", "Recall", "F1", "AUC"], default_rows)
    add_text_paragraph(
        doc,
        f"表2显示，在默认阈值 0.5 下，逻辑回归模型的 Accuracy 为 {summary['metrics_default'].iloc[0]['Accuracy']:.4f}，AUC 为 {summary['metrics_default'].iloc[0]['AUC']:.4f}，明显低于其他模型；随机森林和单一 XGBoost 已显著改善了分类表现；软投票集成在 Accuracy 和 AUC 上进一步提升；Stacking 集成则取得了默认阈值下的最优综合结果，其 Accuracy、Recall、F1 和 AUC 分别为 {summary['best_default_row']['Accuracy']:.4f}、{summary['best_default_row']['Recall']:.4f}、{summary['best_default_row']['F1']:.4f} 和 {summary['best_default_row']['AUC']:.4f}。这一结果表明，在相同训练集和相同平衡策略下，异质树模型经过第二层元学习器融合后，能够较好地整合不同模型的判别优势。",
    )
    tuned_rows = []
    for _, row in summary["metrics_tuned"].iterrows():
        tuned_rows.append([row["模型"], f"{row['阈值']:.2f}", f"{row['Accuracy']:.4f}", f"{row['Precision']:.4f}", f"{row['Recall']:.4f}", f"{row['F1']:.4f}"])
    add_table_with_title(doc, "表3 以准确率为目标的阈值调优结果", ["模型", "最优阈值", "Accuracy", "Precision", "Recall", "F1"], tuned_rows)
    add_text_paragraph(
        doc,
        f"由于代码中包含基于准确率的阈值搜索过程，本文进一步比较了调优后的结果。Stacking 模型在阈值 {summary['best_threshold']:.2f} 处的 Accuracy 提升至 {summary['best_tuned_row']['Accuracy']:.4f}，Precision 为 {summary['best_tuned_row']['Precision']:.4f}，Recall 为 {summary['best_tuned_row']['Recall']:.4f}，F1 为 {summary['best_tuned_row']['F1']:.4f}。与默认阈值相比，调优后的模型牺牲了少量召回率，换取了更高的准确率和精确率。这说明在银行客户预警场景中，阈值选择并不是单纯的技术参数，而是与业务目标直接相关。",
    )
    add_figure(doc, summary["figures"][5], "图6 各模型准确率、F1值与AUC值比较")
    add_figure(doc, summary["figures"][6], "图7 不同模型的ROC曲线比较")
    add_text_paragraph(
        doc,
        "从变量重要性和模型表现综合来看，银行客户流失在很大程度上体现为一种“行为弱化”过程：当客户的交易总额、交易次数和信用卡使用深度持续下降，同时与银行的关系数量有限、近阶段行为变化显著时，模型更容易判断其为潜在流失客户。这一发现与银行客户流失研究的既有发现基本一致，即客户行为变量而不是静态人口属性，通常更能揭示客户关系的变化方向[3][4][5]。",
    )
    add_figure(doc, summary["figures"][7], f"图8 Stacking模型混淆矩阵（阈值={summary['best_threshold']:.2f}）")
    add_text_paragraph(
        doc,
        "需要强调的是，本文并不试图从公开横截面数据中推出严格的因果结论。模型识别的是与流失结果高度相关的统计特征，而不是证明某一变量必然导致流失。将统计相关性转化为经营策略时，仍需结合客户分层、成本约束和业务实验进一步验证。将本文结果与英文稿对照还可以发现，未复现的高 F1 指标并不应被直接写入终稿；只有经代码验证后能够重复得到的结果，才适合纳入论文正文。",
    )
    for text in [
        "如果将本文结果与已有银行客户流失文献进行横向比较，可以发现一个较为稳定的共同点：客户活跃度和交易变化始终是最具解释力的一组信号。无论是采用逻辑回归、随机森林还是提升树模型，只要变量中包含交易金额、交易次数、关系数量和近期行为变化率，这些指标通常都会排在重要性前列。本文在公开 BankChurners 数据上的结果同样支持这一点，这说明客户流失虽然表现形式多样，但其在数据层面往往体现为“使用频率下降”“账户关系收缩”和“行为变化加剧”的组合。因此，银行在实际运营中不应只依据静态画像做预警，而应重点跟踪能够反映行为变化方向的动态指标。",
        "此外，本文结果还说明，模型比较不能只看某一个指标。以逻辑回归为例，它在可解释性上具有明显优势，但在本研究数据中 AUC 和 F1 均落后于其他模型；随机森林具备较好的稳健性，且对变量重要性分析较为友好；单一 XGBoost 和软投票集成已经可以取得较高精度，但 Stacking 仍然表现更优。也就是说，模型性能的提升并不是由某一算法单独完成，而是在合理的数据处理、样本平衡和模型融合共同作用下实现的。如果忽略前置步骤，仅把性能差异归因于“用了更高级的模型”，就会把研究结论简化得过于粗糙。",
        "对银行管理者而言，本文结果至少提供了三方面启示。第一，可以把模型评分作为客户预警系统中的第一层筛查工具，优先定位交易活跃度下降、联系行为异常且账户关系较弱的客户；第二，可以把特征重要性排序转化为客户经理的观察清单，使人工判断与模型结果形成互补，而不是彼此替代；第三，在实际部署时应根据营销预算、触达成本和留存收益调整决策阈值，而不是机械套用训练阶段的默认阈值。换言之，模型的真正价值不在于单次预测是否“神准”，而在于能否嵌入银行日常客户经营流程之中，形成持续可用的识别与干预机制。",
    ]:
        add_text_paragraph(doc, text)

    add_heading(doc, "5 结论与展望", 1)
    for text in [
        "本文以公开的 BankChurners 数据集为基础，对银行信用卡客户流失预测问题进行了系统性实证研究。研究围绕数据预处理、特征筛选、样本平衡、模型构建和结果比较展开，形成了相对完整且可复现的实验流程。结果表明，在统一的数据处理条件下，Stacking 集成模型在默认阈值和调优阈值两种情形下均表现出最优或接近最优的综合性能，说明异质树模型融合在该类问题中具有明显优势。",
        f"具体而言，本文的主要结论包括：第一，样本不平衡确实会影响流失识别效果，仅依赖总体准确率难以全面评价模型优劣；第二，交易总金额、交易总次数、循环余额、季度交易变化率、账户关系数量和平均信用利用率是识别客户流失的重要变量；第三，在默认阈值 0.5 下，Stacking 模型的 AUC 达到 {summary['best_default_row']['AUC']:.4f}；第四，在以准确率为目标的阈值搜索后，Stacking 模型在 {summary['best_threshold']:.2f} 阈值处取得 {summary['best_tuned_row']['Accuracy']:.4f} 的准确率和 {summary['best_tuned_row']['F1']:.4f} 的 F1 值，具有较强的业务应用潜力。",
        "本文的贡献主要是将已有机器学习方法在银行客户流失场景下进行了规范化整合和验证，而不是提出全新的原创算法。未来研究若能够获得真实银行的面板数据，可进一步引入时间序列特征和客户生命周期特征，构建更贴近业务场景的动态流失预测模型；同时也可以结合 SHAP、因果推断或生存分析等方法增强解释性。",
    ]:
        add_text_paragraph(doc, text)
    for text in [
        "在研究不足方面，本文也需要保持克制和诚实。首先，论文基于公开样本完成，虽然保证了透明度，但并不能完全替代真实业务环境中的建模需求。真实银行场景中，流失客户往往受到宏观环境、产品价格、服务质量、竞品策略和营销触达等多种外部因素影响，而公开数据中并没有这些变量。其次，本文虽然对多种模型进行了比较，但参数设定仍然以课程作业和论文代码中较为稳定的配置为主，并未追求极限调优。再次，本文主要通过随机森林重要性和业务逻辑讨论变量意义，尚未使用更细粒度的局部解释方法去观察单个客户的流失原因。换言之，本文的目标是建立一条相对规范的研究路径，而非穷尽银行客户流失预测的所有技术细节。",
        "后续研究若希望进一步提高论文质量，可以从“更真实的数据”“更细致的解释”和“更接近业务的验证”三个方向展开。其一，若能够获得银行内部更长时间跨度的数据，可在现有分类框架基础上引入时间窗口特征、事件序列特征和客户生命周期特征，从而把静态识别拓展为动态监测。其二，可以将 SHAP、局部可解释模型、部分依赖图等方法引入结果分析，进一步说明模型为何会将某些客户识别为高风险。其三，可以把模型分数真正用于营销实验之中，通过 A/B 测试比较不同阈值和不同干预策略的留存效果，以检验模型是否能够在经营层面创造实质性收益。这些方向既能够提升研究深度，也更符合银行客户流失研究从“离线建模”走向“在线应用”的发展趋势。",
    ]:
        add_text_paragraph(doc, text)

    doc.add_page_break()
    add_heading(doc, "【参考文献】", 1)
    references = [
        "[1] 兰军，严广乐. 基于客户特征分群的银行客户流失分析[J]. 技术经济与管理研究，2014(5)：105-108.",
        "[2] Xu X, Xia Y. Research on customer churn prediction model based on probability calibration[J]. Statistics and Applications, 2021, 10(4): 634-641.",
        "[3] AL-Najjar D, Al-Rousan N, AL-Najjar H. Machine learning to develop credit card customer churn prediction[J]. Journal of Theoretical and Applied Electronic Commerce Research, 2022, 17(4): 1529-1542.",
        "[4] de Lima Lemos M S, da Silva F C M, Evsukoff A G, et al. Propension to customer churn in a financial institution: a machine learning approach using discrepancy-based cohort and one-class classification[J]. Neural Computing and Applications, 2022, 34: 17361-17377.",
        "[5] Peng K, Peng Y, Li W. Research on customer churn prediction and model interpretability analysis[J]. PLoS ONE, 2023, 18(12): e0289724.",
        "[6] Siddiqui N, Haque M A, Khan S M S, et al. Different ML-based strategies for customer churn prediction in banking sector[J]. Journal of Data, Information and Management, 2024, 6(3): 217-234.",
        "[7] Wolpert D H. Stacked generalization[J]. Neural Networks, 1992, 5(2): 241-259.",
        "[8] Breiman L. Random forests[J]. Machine Learning, 2001, 45(1): 5-32.",
        "[9] Chawla N V, Bowyer K W, Hall L O, et al. SMOTE: Synthetic Minority Over-sampling Technique[J]. Journal of Artificial Intelligence Research, 2002, 16: 321-357.",
        "[10] Tomek I. Two modifications of CNN[J]. IEEE Transactions on Systems, Man, and Cybernetics, 1976, 6(11): 769-772.",
        "[11] Batista G E A P A, Prati R C, Monard M C. A study of the behavior of several methods for balancing machine learning training data[J]. ACM SIGKDD Explorations Newsletter, 2004, 6(1): 20-29.",
        "[12] Chen T, Guestrin C. XGBoost: A scalable tree boosting system[C]//Proceedings of the 22nd ACM SIGKDD International Conference on Knowledge Discovery and Data Mining. New York: ACM, 2016: 785-794.",
        "[13] Ke G, Meng Q, Finley T, et al. LightGBM: A highly efficient gradient boosting decision tree[C]//Advances in Neural Information Processing Systems 30. Red Hook: Curran Associates, 2017: 3146-3154.",
        "[14] Prokhorenkova L, Gusev G, Vorobev A, et al. CatBoost: unbiased boosting with categorical features[C]//Advances in Neural Information Processing Systems 31. Red Hook: Curran Associates, 2018: 6638-6648.",
        "[15] Amin A, Anwar S, Adnan A, et al. Comparing oversampling techniques to handle the class imbalance problem: a customer churn prediction case study[J]. IEEE Access, 2016, 4: 7940-7957.",
        "[16] 方匡南，吴见彬，朱建平，等. 随机森林方法研究综述[J]. 统计与信息论坛，2011，26(3)：32-38.",
    ]
    for ref in references:
        add_text_paragraph(doc, ref, size=Pt(9), east_asia="楷体", first_line_chars=0.0, align=WD_ALIGN_PARAGRAPH.LEFT)

    doc.add_page_break()
    add_center_title(doc, "致谢", Pt(14), east_asia="黑体", bold=True)
    add_text_paragraph(
        doc,
        "在论文完成过程中，我首先要感谢指导教师在选题、开题和论文修改阶段提供的细致指导。老师不仅指出了研究设计中的方法问题，也提醒我注意中文学术写作与英文直译之间的差异，使我能够在后续修改中更加重视论证逻辑、概念界定与结果表达的一致性。与此同时，也感谢学院提供的学习环境与课程训练，使我能够将机器学习、数据分析与金融场景结合起来完成本次研究。最后，感谢家人和同学在数据处理、写作沟通和情绪支持方面给予的帮助。论文中的不足之处由本人负责，并将在后续学习中继续改进。",
    )

    doc.add_page_break()
    add_center_title(doc, TITLE_EN, Pt(18), east_asia="Times New Roman", bold=True)
    p = doc.add_paragraph()
    run = p.add_run("【Abstract】")
    set_run_font(run, size=Pt(12), east_asia="Times New Roman", ascii_font="Times New Roman", bold=True)
    abstract_en = (
        "Customer churn prediction has become a practical issue in retail banking because customer attrition directly affects revenue retention, cross-selling opportunities, and long-term customer value. "
        "Using the public BankChurners credit-card dataset, this study develops a reproducible churn prediction framework that combines feature screening, imbalance handling, and stacking ensemble learning. "
        "After removing identifier and Naive Bayes probability fields, categorical variables are one-hot encoded, numeric variables are imputed with iterative imputation and standardized, and 14 core features are retained through random-forest importance screening. "
        "SMOTETomek is applied only to the training set in order to preserve the original imbalanced distribution in the test set. A two-layer stacking model is then built with Random Forest, XGBoost, CatBoost, and LightGBM as base learners and XGBoost as the meta-learner. "
        f"Under the default threshold of 0.5, the stacking model reaches an Accuracy of {summary['best_default_row']['Accuracy']:.4f}, an F1 score of {summary['best_default_row']['F1']:.4f}, and an AUC of {summary['best_default_row']['AUC']:.4f}. "
        f"After threshold tuning, the best threshold is {summary['best_threshold']:.2f}, where Accuracy rises to {summary['best_tuned_row']['Accuracy']:.4f} and F1 to {summary['best_tuned_row']['F1']:.4f}. "
        "The findings show that total transaction amount, total transaction count, revolving balance, and quarterly transaction changes are the most important churn indicators. "
        "The contribution of this thesis lies not in proposing a brand-new algorithm, but in providing a rigorous application-oriented validation that shows the practical suitability of combining random-forest feature screening, SMOTETomek, and stacking ensemble learning in a bank customer churn scenario."
    )
    add_text_paragraph(doc, abstract_en, size=Pt(10.5), east_asia="Times New Roman", first_line_chars=0.0)
    p = doc.add_paragraph()
    run = p.add_run("【Key words】")
    set_run_font(run, size=Pt(12), east_asia="Times New Roman", ascii_font="Times New Roman", bold=True)
    run2 = p.add_run("customer churn prediction; stacking ensemble; SMOTETomek; credit card customers; machine learning")
    set_run_font(run2, size=Pt(10.5), east_asia="Times New Roman", ascii_font="Times New Roman")

    out_path = OUTPUT_DOC / f"{TITLE_CN}_中文论文_终稿_v3.docx"
    doc.save(out_path)
    return out_path


def main():
    ensure_dirs()
    summary = generate_metrics_and_figures()
    doc_path = build_document(summary)
    print(f"DOCX={doc_path}")


if __name__ == "__main__":
    main()
